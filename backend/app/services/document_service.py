"""
Document processing service.
Extracts text from PDF, DOCX, and TXT files.
"""
import io
from fastapi import UploadFile


def extract_text_sync(content: bytes, filename: str) -> str:
    """Extract plain text from raw file bytes based on filename extension."""
    fname = filename.lower()

    if fname.endswith(".txt"):
        return content.decode("utf-8", errors="ignore")

    if fname.endswith(".pdf"):
        return _extract_pdf(content)

    if fname.endswith(".docx"):
        return _extract_docx(content)

    raise ValueError(f"Unsupported file type: {filename}")


async def extract_text(file: UploadFile) -> str:
    """Extract plain text from uploaded PDF, DOCX, or TXT file (async version)."""
    content = await file.read()
    return extract_text_sync(content, file.filename)


def _extract_pdf(content: bytes) -> str:
    """Extract text from PDF using pypdf."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)
    except ImportError:
        raise RuntimeError("pypdf not installed. Run: pip install pypdf")
    except Exception as e:
        raise RuntimeError(f"PDF extraction failed: {e}")


def _extract_docx(content: bytes) -> str:
    """Extract text from DOCX using python-docx — includes paragraphs AND table cells."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(content))
        lines = []

        # Paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                lines.append(p.text.strip())

        # Tables — critical for menu documents where items are in table rows
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))

        return "\n".join(lines)
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        raise RuntimeError(f"DOCX extraction failed: {e}")
